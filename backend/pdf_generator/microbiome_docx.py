"""
Microbiome DOCX Report Generator
Generates a polished Word (.docx) document from microbiome XLSX/CSV analysis data.

Mirrors the structure of microbiome_pdf.py but uses python-docx for native Word output:
  - Cover page with patient info and sections list
  - Per-section pages: title, description, subsection headers, data tables
  - Summary page (alarmed parameters grouped by category)
  - Doctor comments appendix (if any)
  - Consistent header/footer with patient name, report N° and page number

Requires: python-docx >= 1.1.0
"""

import io
import os
import re
from collections import OrderedDict

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Shared constants — imported directly to avoid re-defining large structures
# (ReportLab is also installed, so the import is safe)
from .microbiome_pdf import SECTION_DESCRIPTIONS, SUBSECTION_MAP

# ── Colour palette ────────────────────────────────────────────────────────────

C_CYAN       = RGBColor(0x16, 0xBA, 0xDE)
C_BLACK      = RGBColor(0x00, 0x00, 0x00)
C_DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
C_MID_GRAY   = RGBColor(0x88, 0x88, 0x88)

# Hex strings for XML cell shading (no '#' prefix)
HEX_LIGHT_GRAY = 'FAFAFA'   # parent-row background
HEX_COVER_BG   = 'ECEBE5'   # cover page lower-section background
HEX_WHITE      = 'FFFFFF'
HEX_COMMENT_BG = 'F9F9F7'   # doctor comment box background

# ── Local alarm map (avoids pulling in ReportLab colour objects) ──────────────

_ALARM_MAP: dict = {
    '_':         ('',  False),
    'Asterisco': ('!', True),
    'A':         ('!', True),
    'AA':        ('!', True),
    'AAAA':      ('!', True),
    'B':         ('!', True),
    'BB':        ('!', True),
    'BBBB':      ('!', True),
    '+/-':       ('!', True),
    'R':         ('!', True),
}


# ── XML / OOXML helper functions ──────────────────────────────────────────────

def _shd(cell, hex_color: str) -> None:
    """Set table-cell background shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    tcPr.append(shd)


def _page_number_field(run) -> None:
    """Insert a {PAGE} auto-field into a run."""
    for tag, ftype in (('w:fldChar', 'begin'), ('w:instrText', None), ('w:fldChar', 'end')):
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = ' PAGE '
        else:
            el.set(qn('w:fldCharType'), ftype)
        run._r.append(el)


def _para_bottom_border(para, color: str, size: int = 6, space: int = 4) -> None:
    """Add a bottom border to a paragraph (used for rules / separators)."""
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), str(space))
    bot.set(qn('w:color'), color.lstrip('#'))
    pBdr.append(bot)
    pPr.append(pBdr)


def _tbl_borders(tbl, *, outer_sz: int = 0, inner_sz: int = 2,
                 inner_color: str = 'E8E8E8') -> None:
    """Configure table borders: no outer borders, very light inner horizontal lines."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none' if outer_sz == 0 else 'single')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    ih = OxmlElement('w:insideH')
    ih.set(qn('w:val'),   'single' if inner_sz > 0 else 'none')
    ih.set(qn('w:sz'),    str(inner_sz))
    ih.set(qn('w:space'), '0')
    ih.set(qn('w:color'), inner_color.lstrip('#'))
    tblBorders.append(ih)
    tblPr.append(tblBorders)


def _cell_bottom_border(cell, color: str = 'BBBBBB', size: int = 4) -> None:
    """Add a thin bottom border to a single cell (table header separator)."""
    tcPr     = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '0')
    bot.set(qn('w:color'), color.lstrip('#'))
    tcBorders.append(bot)
    tcPr.append(tcBorders)


def _tbl_full_width(tbl) -> None:
    """Set a table to 100 % page width."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def _right_tab(para, pos_twips: int = 9355) -> None:
    """Add a right-aligned tab stop to a paragraph (for header/footer)."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


# ── Main generator ────────────────────────────────────────────────────────────

class MicrobiomeDOCXGenerator:
    """
    Build a polished A4 Word document from a microbiome analysis DataFrame.

    Usage:
        docx_bytes = MicrobiomeDOCXGenerator(df).generate()
    """

    # A4: 21 cm wide, 2 × 1.6 cm margins → 17.8 cm available
    AVAIL_CM = 21 - 1.6 - 1.6   # ≈ 17.8 cm

    # Column widths as fractions of available width
    COL_FRACTIONS = [0.40, 0.12, 0.12, 0.22, 0.14]

    def __init__(self, df: pd.DataFrame,
                 comments: dict | None = None,
                 cited_params: set | None = None):
        self.df           = df.copy()
        self.comments     = {int(k): str(v) for k, v in (comments or {}).items()}
        self.cited_params = set(cited_params) if cited_params else set()

        # Resolve logo path — works both locally and inside Docker
        _base        = os.path.dirname(os.path.abspath(__file__))
        _repo        = os.path.normpath(os.path.join(_base, '..', '..'))
        _logo_local  = os.path.join(_repo, 'frontend', 'logo_bw.png')
        self._logo_path: str | None = (
            _logo_local if os.path.exists(_logo_local)
            else ('/app/logo_bw.png' if os.path.exists('/app/logo_bw.png') else None)
        )

        self._extract_patient_info()

    # ── Patient info ──────────────────────────────────────────────────────────

    def _extract_patient_info(self) -> None:
        r = self.df.iloc[0]

        def g(col: str, default: str = '') -> str:
            v = r.get(col, default)
            return '' if pd.isna(v) else str(v).strip()

        self.patient_name  = g('DescripcionMuestra', 'Unknown patient')
        self.patient_id    = g('DNI')
        self.client        = g('Cliente')
        self.report_number = g('NumInforme')
        self.sample_date   = self._fmt_date(g('FechaMuestra'))
        self.valid_date    = self._fmt_date(g('Validacion'))

    @staticmethod
    def _fmt_date(s: str) -> str:
        if not s:
            return ''
        try:
            return pd.to_datetime(s, dayfirst=True).strftime('%d/%m/%Y')
        except Exception:
            return s

    # ── Data helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def _strip_html(text) -> str:
        if not text or str(text).strip() in ('', 'nan', 'None'):
            return ''
        t = re.sub(r'<[^>]+>', ' ', str(text))
        return re.sub(r'\s+', ' ', t).strip()

    @staticmethod
    def _clean(val) -> str:
        s = str(val).strip()
        return '' if s in ('', 'nan', 'None') else s

    @staticmethod
    def _clean_param(val) -> str:
        """Strip trailing [CODE] identifiers from parameter names."""
        s = str(val).strip()
        if s in ('', 'nan', 'None'):
            return ''
        return re.sub(r'\s*\[[^\[\]]+\]\s*$', '', s).strip()

    def _result(self, row) -> str:
        r1 = self._clean(row.get('Resultado1', ''))
        r2 = self._clean(row.get('Resultado2', ''))
        if r1:
            return r1
        if r2:
            return r2[:40] + ('…' if len(r2) > 40 else '')
        return '—'

    def _ref_range(self, row) -> str:
        vmax = self._clean(row.get('VRMaximo', '')).replace(',', '.')
        vmin = self._clean(row.get('VRMinimo', '')).replace(',', '.')
        if vmin and vmax:
            return f'{vmin} – {vmax}'
        if vmax:
            return f'< {vmax}'
        if vmin:
            return f'> {vmin}'
        return '—'

    def _alarm(self, row) -> tuple:
        """Return (symbol: str, alarmed: bool)."""
        if str(row.get('Alarma', 'Falso')).strip() != 'Verdadero':
            return '', False
        code = str(row.get('AlarmaDescripcion', '_')).strip()
        return _ALARM_MAP.get(code, ('!', True))

    # ── Document / page setup ─────────────────────────────────────────────────

    def _new_document(self) -> Document:
        doc = Document()

        # A4 page size + margins
        sec = doc.sections[0]
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(1.6)
        sec.right_margin  = Cm(1.6)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(3.2)

        # Remove default paragraph spacing from Normal style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after  = Pt(0)

        return doc

    def _setup_header_footer(self, doc: Document) -> None:
        """
        Configure header + footer for the main content pages.
        Cover page uses a different (blank) first-page header.
        """
        sec = doc.sections[0]
        sec.different_first_page_header_footer = True  # cover = no header

        # ── Running header ────────────────────────────────────────────────────
        hdr_p = sec.header.paragraphs[0]
        hdr_p.clear()
        hdr_p.paragraph_format.space_before = Pt(0)
        hdr_p.paragraph_format.space_after  = Pt(0)
        _right_tab(hdr_p)

        rl = hdr_p.add_run(self.patient_name)
        rl.font.name = 'Calibri'
        rl.font.size = Pt(8)
        rl.font.color.rgb = C_DARK_GRAY

        hdr_p.add_run('\t')

        rr = hdr_p.add_run(f'Report N° {self.report_number}')
        rr.font.name = 'Calibri'
        rr.font.size = Pt(7.5)
        rr.font.color.rgb = C_MID_GRAY

        # Cyan bottom border on the header
        _para_bottom_border(hdr_p, color='16BADE', size=6, space=4)

        # ── Footer ────────────────────────────────────────────────────────────
        ftr_p = sec.footer.paragraphs[0]
        ftr_p.clear()
        ftr_p.paragraph_format.space_before = Pt(0)
        ftr_p.paragraph_format.space_after  = Pt(0)
        _right_tab(ftr_p)

        rc = ftr_p.add_run(self.client)
        rc.font.name = 'Calibri'
        rc.font.size = Pt(7.5)
        rc.font.color.rgb = C_MID_GRAY

        ftr_p.add_run('\t')

        rpg = ftr_p.add_run()
        rpg.font.name = 'Calibri'
        rpg.font.size = Pt(9)
        rpg.font.bold = True
        rpg.font.color.rgb = C_DARK_GRAY
        _page_number_field(rpg)

    # ── Cover page ────────────────────────────────────────────────────────────

    def _build_cover(self, doc: Document) -> None:
        """Cover page: logo (optional), title, patient info band at 0.38 of page height."""

        # ── Layout constants ──────────────────────────────────────────────────
        # A4: 841.89 pt high, top margin 1.6 cm = 45.35 pt
        # Band target: 0.38 × 841.89 − 45.35 ≈ 274.5 pt from content start
        _LOGO_W          = Cm(7.5)           # logo display width
        _LOGO_W_PT       = 7.5 * 28.35       # ≈ 212.6 pt
        _LOGO_ASPECT     = 344 / 85          # px ratio of logo_bw.png
        _LOGO_H_PT       = _LOGO_W_PT / _LOGO_ASPECT   # ≈ 52.5 pt
        _LOGO_BEFORE     = 50.0              # pt above logo
        _LOGO_AFTER      = 10.0
        _TITLE_BEFORE    = 10.0
        _TITLE_H         = 26.0              # approx rendered height for Pt(22)
        _TITLE_AFTER     = 6.0
        _content_used    = (_LOGO_BEFORE + _LOGO_H_PT + _LOGO_AFTER
                            + _TITLE_BEFORE + _TITLE_H + _TITLE_AFTER)
        _band_target_pt  = 0.38 * 841.89 - 45.35        # ≈ 274.5 pt
        _spacer_pt       = max(_band_target_pt - _content_used, 0)

        # ── Optional logo ─────────────────────────────────────────────────────
        logo_drawn = False
        if self._logo_path:
            try:
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.paragraph_format.space_before = Pt(_LOGO_BEFORE)
                p_logo.paragraph_format.space_after  = Pt(_LOGO_AFTER)
                p_logo.add_run().add_picture(self._logo_path, width=_LOGO_W)
                logo_drawn = True
            except Exception:
                pass

        # ── Title ─────────────────────────────────────────────────────────────
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_before = Pt(_TITLE_BEFORE if logo_drawn else 120)
        p_t.paragraph_format.space_after  = Pt(_TITLE_AFTER)
        rt = p_t.add_run('MICROBIOME ANALYSIS')
        rt.font.name  = 'Calibri Light'
        rt.font.size  = Pt(22)
        rt.font.color.rgb = C_BLACK

        # ── Spacer — pushes band to 0.38 of page height ───────────────────────
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(_spacer_pt if logo_drawn else 18)
        p_sp.paragraph_format.space_after  = Pt(0)

        # ── Patient info (2-column table, cover-bg shading) ───────────────────
        rows_data = [
            ('Patient',    self.patient_name,  'Collection', self.sample_date),
            ('ID',         self.patient_id,    'Validation', self.valid_date),
            ('Clinic',     self.client,        'Report N°',  self.report_number),
        ]
        info_tbl = doc.add_table(rows=len(rows_data), cols=2)
        info_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _tbl_full_width(info_tbl)
        # Remove all borders from this info table
        _tbl_borders(info_tbl, outer_sz=0, inner_sz=0)

        for r_idx, (l1, v1, l2, v2) in enumerate(rows_data):
            row = info_tbl.rows[r_idx]
            for cell in row.cells:
                _shd(cell, HEX_COVER_BG)

            for cell, label, value in (
                (row.cells[0], l1, v1),
                (row.cells[1], l2, v2),
            ):
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(4)

                run_lbl = para.add_run(f'{label}:  ')
                run_lbl.font.name  = 'Calibri'
                run_lbl.font.size  = Pt(9)
                run_lbl.font.bold  = True
                run_lbl.font.color.rgb = C_DARK_GRAY

                run_val = para.add_run(value)
                run_val.font.name  = 'Calibri'
                run_val.font.size  = Pt(9)
                run_val.font.color.rgb = C_BLACK

        # ── Thin separator ────────────────────────────────────────────────────
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(16)
        p_sep.paragraph_format.space_after  = Pt(6)
        _para_bottom_border(p_sep, color='DDDDDD', size=4, space=1)

        # ── Sections included ─────────────────────────────────────────────────
        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after  = Pt(6)
        rl = p_lbl.add_run('Sections included:')
        rl.font.name  = 'Calibri'
        rl.font.size  = Pt(9)
        rl.font.bold  = True
        rl.font.color.rgb = C_DARK_GRAY

        for sec_name in self.df['TipoInforme'].unique():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent   = Cm(0.5)
            p.paragraph_format.space_before  = Pt(0)
            p.paragraph_format.space_after   = Pt(2)
            r = p.add_run(f'• {sec_name}')
            r.font.name  = 'Calibri'
            r.font.size  = Pt(9)
            r.font.color.rgb = C_DARK_GRAY

        doc.add_page_break()

    # ── Data table ────────────────────────────────────────────────────────────

    def _build_data_table(self, doc: Document, sec_df: pd.DataFrame,
                          notes_out: list | None = None) -> None:
        """Build the 5-column Parameter / Result / Unit / Reference Range / Status table."""
        col_widths = [Cm(self.AVAIL_CM * f) for f in self.COL_FRACTIONS]
        hdrs       = ['Parameter', 'Result', 'Unit', 'Reference Range', '']
        hdr_align  = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4

        tbl = doc.add_table(rows=1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _tbl_full_width(tbl)
        _tbl_borders(tbl, outer_sz=0, inner_sz=2, inner_color='E8E8E8')

        # Set column widths via the first row's cells
        for col_idx, width in enumerate(col_widths):
            for cell in tbl.column_cells(col_idx):
                cell.width = width

        # ── Header row ────────────────────────────────────────────────────────
        hdr_row = tbl.rows[0]
        for cell, label, align in zip(hdr_row.cells, hdrs, hdr_align):
            _shd(cell, HEX_WHITE)
            _cell_bottom_border(cell)
            para = cell.paragraphs[0]
            para.alignment = align
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after  = Pt(5)
            run = para.add_run(label)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(8.5)
            run.font.bold  = True
            run.font.color.rgb = C_DARK_GRAY

        # ── Data rows ─────────────────────────────────────────────────────────
        for _, row in sec_df.iterrows():
            param    = self._clean_param(row.get('Ensayo', ''))
            is_child = param.startswith('- ')
            param_key = param.lstrip('- ').strip()
            result   = self._result(row)
            unit     = self._clean(row.get('Unidad1', ''))
            ref      = self._ref_range(row)
            sym, alarmed = self._alarm(row)

            # Doctor-comment citation mark
            cited_mark  = ' ✉' if param_key in self.cited_params else ''
            sym_display = (sym + cited_mark) if sym else cited_mark

            # Memo note
            memo = self._strip_html(row.get('Memo', ''))
            if memo and notes_out is not None:
                notes_out.append(f'[{param_key}]  {memo}')

            dr = tbl.add_row()

            # Parent rows: light gray background
            if not is_child:
                for cell in dr.cells:
                    _shd(cell, HEX_LIGHT_GRAY)

            # ── Cell 0: Parameter ─────────────────────────────────────────────
            c0   = dr.cells[0]
            p0   = c0.paragraphs[0]
            p0.paragraph_format.left_indent  = Cm(0.4) if is_child else Cm(0.1)
            p0.paragraph_format.space_before = Pt(3)
            p0.paragraph_format.space_after  = Pt(3)
            r0 = p0.add_run(param_key)
            r0.font.name  = 'Calibri'
            r0.font.size  = Pt(8.5)
            r0.font.bold  = not is_child
            r0.font.color.rgb = C_DARK_GRAY if is_child else C_BLACK

            # ── Cells 1-4: Result, Unit, Ref, Status ─────────────────────────
            for ci, (content, align) in enumerate([
                (result,      WD_ALIGN_PARAGRAPH.CENTER),
                (unit,        WD_ALIGN_PARAGRAPH.CENTER),
                (ref,         WD_ALIGN_PARAGRAPH.CENTER),
                (sym_display, WD_ALIGN_PARAGRAPH.CENTER),
            ], start=1):
                para = dr.cells[ci].paragraphs[0]
                para.alignment = align
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                run = para.add_run(content)
                run.font.name  = 'Calibri'
                run.font.size  = Pt(7) if len(content) > 14 else Pt(8.5)
                run.font.color.rgb = C_DARK_GRAY

    # ── Section builder ───────────────────────────────────────────────────────

    def _build_section(self, doc: Document, section: str) -> None:
        """Build one analysis section with subsection headers and data tables."""
        sec_df = self.df[self.df['TipoInforme'] == section]

        # Section title (cyan, large)
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after  = Pt(6)
        rt = p_t.add_run(section)
        rt.font.name  = 'Calibri Light'
        rt.font.size  = Pt(20)
        rt.font.color.rgb = C_CYAN

        # Section description
        desc = SECTION_DESCRIPTIONS.get(section, '')
        if desc:
            p_d = doc.add_paragraph()
            p_d.paragraph_format.space_before = Pt(0)
            p_d.paragraph_format.space_after  = Pt(10)
            rd = p_d.add_run(desc)
            rd.font.name  = 'Calibri'
            rd.font.size  = Pt(8.5)
            rd.font.color.rgb = C_DARK_GRAY

        # ── Subsection splitting (mirrors PDF logic) ──────────────────────────
        subsections = SUBSECTION_MAP.get(section, [])

        if subsections:
            trigger_to_idx: dict = {}
            for sub_idx, (title, sdesc, triggers) in enumerate(subsections):
                if isinstance(triggers, str):
                    triggers = [triggers]
                for t in triggers:
                    trigger_to_idx[t] = sub_idx

            sub_meta = {i: (title, sdesc) for i, (title, sdesc, _) in enumerate(subsections)}

            pre_rows:        list = []
            row_assignments: list = []
            current_idx           = None

            for _, row in sec_df.iterrows():
                cleaned = self._clean_param(row.get('Ensayo', ''))
                if cleaned in trigger_to_idx:
                    current_idx = trigger_to_idx[cleaned]
                if current_idx is None:
                    pre_rows.append(row.name)
                else:
                    row_assignments.append((current_idx, row.name))

            merged: OrderedDict = OrderedDict()
            for sub_idx, ridx in row_assignments:
                merged.setdefault(sub_idx, []).append(ridx)
            if pre_rows and merged:
                first_key = next(iter(merged))
                merged[first_key] = pre_rows + merged[first_key]

            buckets = [(*sub_meta[sidx], indices)
                       for sidx, indices in sorted(merged.items())]

            for sub_title, sub_desc, row_idx in buckets:
                sub_df = sec_df.loc[row_idx]
                if sub_df.empty:
                    continue

                # Subsection title
                p_s = doc.add_paragraph()
                p_s.paragraph_format.space_before = Pt(14)
                p_s.paragraph_format.space_after  = Pt(2)
                rs = p_s.add_run(sub_title)
                rs.font.name  = 'Calibri'
                rs.font.size  = Pt(10)
                rs.font.bold  = True
                rs.font.color.rgb = C_DARK_GRAY

                # Subsection description
                if sub_desc:
                    p_sd = doc.add_paragraph()
                    p_sd.paragraph_format.space_before = Pt(0)
                    p_sd.paragraph_format.space_after  = Pt(6)
                    rsd = p_sd.add_run(sub_desc)
                    rsd.font.name   = 'Calibri'
                    rsd.font.size   = Pt(8)
                    rsd.font.italic = True
                    rsd.font.color.rgb = C_MID_GRAY

                notes: list = []
                self._build_data_table(doc, sub_df, notes_out=notes)

                for note in notes:
                    pn = doc.add_paragraph()
                    pn.paragraph_format.left_indent  = Cm(0.5)
                    pn.paragraph_format.space_before = Pt(2)
                    pn.paragraph_format.space_after  = Pt(1)
                    rn = pn.add_run(note)
                    rn.font.name   = 'Calibri'
                    rn.font.size   = Pt(7.5)
                    rn.font.italic = True
                    rn.font.color.rgb = C_MID_GRAY

                # Breathing space after each subsection table
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)

        else:
            # No subsections — single flat table
            notes: list = []
            self._build_data_table(doc, sec_df, notes_out=notes)
            for note in notes:
                pn = doc.add_paragraph()
                pn.paragraph_format.left_indent  = Cm(0.5)
                pn.paragraph_format.space_before = Pt(2)
                rn = pn.add_run(note)
                rn.font.name   = 'Calibri'
                rn.font.size   = Pt(7.5)
                rn.font.italic = True
                rn.font.color.rgb = C_MID_GRAY

    # ── Summary section ───────────────────────────────────────────────────────

    def _build_summary(self, doc: Document) -> None:
        """Summary page: alarmed parameters grouped by category/subsection."""

        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after  = Pt(6)
        rt = p_t.add_run('Summary')
        rt.font.name  = 'Calibri Light'
        rt.font.size  = Pt(20)
        rt.font.color.rgb = C_CYAN

        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_after = Pt(10)
        rd = p_d.add_run(
            'The following parameters were found outside their reference ranges. '
            'They are grouped by analysis category and section for a quick clinical overview.'
        )
        rd.font.name  = 'Calibri'
        rd.font.size  = Pt(8.5)
        rd.font.color.rgb = C_DARK_GRAY

        # ── Collect alarmed rows ──────────────────────────────────────────────
        entries: list = []
        for section in self.df['TipoInforme'].unique():
            sec_df     = self.df[self.df['TipoInforme'] == section]
            subsections = SUBSECTION_MAP.get(section, [])

            if subsections:
                trigger_to_idx: dict = {}
                for sub_idx, (title, _, triggers) in enumerate(subsections):
                    if isinstance(triggers, str):
                        triggers = [triggers]
                    for t in triggers:
                        trigger_to_idx[t] = sub_idx

                sub_meta = {i: title for i, (title, _, _) in enumerate(subsections)}

                pre_rows:        list = []
                row_assignments: list = []
                current_idx           = None

                for _, row in sec_df.iterrows():
                    cleaned = self._clean_param(row.get('Ensayo', ''))
                    if cleaned in trigger_to_idx:
                        current_idx = trigger_to_idx[cleaned]
                    if current_idx is None:
                        pre_rows.append(row.name)
                    else:
                        row_assignments.append((current_idx, row.name))

                merged: OrderedDict = OrderedDict()
                for sub_idx, ridx in row_assignments:
                    merged.setdefault(sub_idx, []).append(ridx)
                if pre_rows and merged:
                    first_key = next(iter(merged))
                    merged[first_key] = pre_rows + merged[first_key]

                for sub_idx, indices in sorted(merged.items()):
                    sub_df  = sec_df.loc[indices]
                    alarmed = [r for _, r in sub_df.iterrows() if self._alarm(r)[0] == '!']
                    if alarmed:
                        entries.append((section, sub_meta.get(sub_idx, ''), alarmed))
            else:
                alarmed = [r for _, r in sec_df.iterrows() if self._alarm(r)[0] == '!']
                if alarmed:
                    entries.append((section, None, alarmed))

        if not entries:
            p = doc.add_paragraph()
            r = p.add_run(
                'No parameters outside reference ranges were detected in this report.')
            r.font.name   = 'Calibri'
            r.font.size   = Pt(9)
            r.font.italic = True
            r.font.color.rgb = C_MID_GRAY
            return

        current_cat: str | None = None
        for category, sub_title, alarmed_rows in entries:

            # Category heading (once per category)
            if category != current_cat:
                current_cat = category
                pc = doc.add_paragraph()
                pc.paragraph_format.space_before = Pt(16)
                pc.paragraph_format.space_after  = Pt(4)
                rc = pc.add_run(category)
                rc.font.name  = 'Calibri Light'
                rc.font.size  = Pt(13)
                rc.font.color.rgb = C_CYAN

            # Subsection label (if any)
            if sub_title:
                ps = doc.add_paragraph()
                ps.paragraph_format.space_before = Pt(8)
                ps.paragraph_format.space_after  = Pt(2)
                rs = ps.add_run(sub_title)
                rs.font.name  = 'Calibri'
                rs.font.size  = Pt(9)
                rs.font.bold  = True
                rs.font.color.rgb = C_DARK_GRAY

            self._build_data_table(doc, pd.DataFrame(alarmed_rows))

            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(6)

    # ── Doctor comments appendix ──────────────────────────────────────────────

    def _build_doctor_comments(self, doc: Document) -> None:
        """Append doctor notes as a structured appendix (page-keyed comments)."""
        if not self.comments:
            return

        doc.add_page_break()

        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(14)
        rt = p_t.add_run('Doctor Comments')
        rt.font.name  = 'Calibri Light'
        rt.font.size  = Pt(18)
        rt.font.color.rgb = C_CYAN

        for page_num in sorted(self.comments):
            text = self.comments[page_num]

            # Shaded box via a 1-cell table
            tbl = doc.add_table(rows=1, cols=1)
            cell = tbl.rows[0].cells[0]
            _shd(cell, HEX_COMMENT_BG)
            _tbl_borders(tbl, outer_sz=4, inner_sz=0)

            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)

            rl = para.add_run(f'PAGE {page_num} — ')
            rl.font.name  = 'Calibri'
            rl.font.size  = Pt(7.5)
            rl.font.bold  = True
            rl.font.color.rgb = C_MID_GRAY

            rc = para.add_run(text)
            rc.font.name  = 'Calibri'
            rc.font.size  = Pt(8.5)
            rc.font.color.rgb = C_DARK_GRAY

            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(8)

    # ── Public API ────────────────────────────────────────────────────────────

    def generate(self) -> bytes:
        """Build the full Word document and return as raw bytes."""
        doc = self._new_document()
        self._setup_header_footer(doc)
        self._build_cover(doc)

        sections = list(self.df['TipoInforme'].unique())
        for idx, section in enumerate(sections):
            self._build_section(doc, section)
            if idx < len(sections) - 1:
                doc.add_page_break()

        doc.add_page_break()
        self._build_summary(doc)
        self._build_doctor_comments(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ── Convenience wrapper ───────────────────────────────────────────────────────

def generate_microbiome_docx(df: pd.DataFrame,
                             comments: dict | None = None,
                             cited_params: set | None = None) -> bytes:
    """Generate a microbiome Word document from *df* and return raw bytes."""
    return MicrobiomeDOCXGenerator(
        df, comments=comments, cited_params=cited_params
    ).generate()
